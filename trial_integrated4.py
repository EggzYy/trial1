import os
os.environ['TF_ENABLE_ONEDNN_OPTS'] ="0"
os.environ['AUTOGEN_USE_DOCKER'] = "False"
import autogen
from typing import List, Tuple, Dict, Callable
import logging
import time
from dataclasses import dataclass
from retry import retry
import tiktoken
from langchain_text_splitters import RecursiveCharacterTextSplitter
from ollama import Client
from tqdm import tqdm
import json
import re
import os
from typing import List, Tuple, Dict, Callable
import logging
import time
from transformers import pipeline
import numpy as np
from collections import OrderedDict
from unstructured.partition.auto import partition
from unstructured.chunking.title import chunk_by_title
import docx
from collections import Counter
import uuid

# First, add these debug statements at the top of your file after the imports
logging.getLogger('autogen').setLevel(logging.DEBUG)
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')

# --- Data Structures ---
@dataclass
class Usage:
    total_tokens: int
    prompt_tokens: int
    completion_tokens: int

    def add(self, other: "Usage") -> "Usage":
        """Add two Usage objects together."""
        return Usage(
            total_tokens=self.total_tokens + other.total_tokens,
            prompt_tokens=self.prompt_tokens + other.prompt_tokens,
            completion_tokens=self.completion_tokens + other.completion_tokens,
        )

    def to_dict(self) -> dict:
        """Convert the Usage object to a dictionary."""
        return self.__dict__

# --- Abstract LLM Provider Class ---
class LLMProvider:
    """
    Abstract base class for Large Language Model (LLM) providers.

    This class defines the interface that all LLM provider implementations
    should adhere to. It includes methods for calculating token counts,
    splitting text, and getting completions from an LLM.
    """
    def __init__(self) -> None:
        """Initialize the LLM provider"""
        self.model: str = ""

    @staticmethod
    def calculate_chunk_size(token_count: int, token_limit: int) -> int:
        """Calculate the chunk size based on the token count and token limit."""
        if token_count <= token_limit:
            return token_count
        num_chunks = (token_count + token_limit - 1) // token_limit
        chunk_size = token_count // num_chunks
        remaining_tokens = token_count % token_limit
        if remaining_tokens > 0:
            chunk_size += remaining_tokens // num_chunks
        return chunk_size

    def num_tokens_in_string(self, input_str: str) -> int:
        """Calculate the number of tokens in a given string."""
        raise NotImplementedError("The num_tokens_in_string method must be implemented in a subclass.")

    def split_text(self, text: str, max_tokens: int = 1024) -> List[str]:
         """Split the input text into chunks of tokens."""
         raise NotImplementedError("The split_text method must be implemented in a subclass.")

    def get_completion(
        self,
        prompt: str,
        system_message: str = "You are a helpful assistant.",
    ) -> Tuple[str, Usage]:
        """Get the completion for a given prompt."""
        raise NotImplementedError("The get_completion method must be implemented in a subclass.")

# --- OpenAI Compatible LLM Provider ---
class OpenAICompatible(LLMProvider):
    """
    A base class for LLM providers that are compatible with the OpenAI API structure,
    particularly for models accessed via a custom client like Ollama's.
    It handles token calculation using tiktoken and provides a common structure
    for splitting text and getting completions.
    """
    def __init__(self, client, base_url=None, **kwargs):
        """
        Initializes the OpenAI compatible client.

        Args:
            client: The client object used to interact with the LLM (e.g., Ollama client).
            base_url (Optional[str]): The base URL of the LLM API.
            **kwargs: Additional keyword arguments, including 'model' for the model name.
        """
        self.client = client
        self.model = kwargs.get("model", "")
        #self.temperature = kwargs.get("temperature", 0.1)
        self.base_url = base_url
        self.verify = False
        self.encoding = tiktoken.get_encoding("cl100k_base")


    def num_tokens_in_string(self, input_str: str) -> int:
        """Calculates the number of tokens in the given string using tiktoken."""
        num_tokens = len(self.encoding.encode(input_str))
        return num_tokens

    def split_text(self, text: str, max_tokens: int = 1024) -> List[str]:
        """Splits the input text into chunks of tokens."""
        num_tokens_in_text = self.num_tokens_in_string(text)
        token_size = min(max_tokens, num_tokens_in_text)
        logging.info(f"Token size for splitting: {token_size}")
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=token_size,
            chunk_overlap=0,
            separators=["\n\n", "\n", "."],
        )
        return text_splitter.split_text(text)

    @retry(tries=3, delay=2)
    def get_completion(
        self,
        prompt: str,
        system_message: str = "You are a helpful assistant.",
        **kwargs,
    ) -> Tuple[str, Usage]:
        """Gets a completion from the LLM with retry logic."""
        model = kwargs.get("model", self.model)
        #temperature = kwargs.get("temperature", self.temperature)

        try:
            prompt_tokens = self.num_tokens_in_string(system_message + prompt)
            response = self.client.generate(
                model=model,
                system=system_message,
                prompt=prompt,
                options={"top_p": 1}, #"temperature": temperature, 
                stream=False,
            )
            response_content = response.get('response', '')
            completion_tokens = self.num_tokens_in_string(response_content)
            usage = Usage(
                total_tokens=prompt_tokens+completion_tokens,
                prompt_tokens=prompt_tokens,
                completion_tokens=completion_tokens
            )
            return response_content, usage
        except Exception as e:
            logging.error(f"Error during completion: {e}")
            return "", Usage(total_tokens=0, prompt_tokens=0, completion_tokens=0) # Return empty string and default Usage object in case of failure

# --- Concrete OpenAI Implementation ---
class OpenAI(OpenAICompatible):
    """Concrete implementation for OpenAI-compatible models."""
    def __init__(self, api_key: str, base_url: str, **kwargs):
        """Initializes the OpenAI client with the API key and base URL."""
        client = Client(host=base_url.replace("/api", ""))
        super().__init__(client, base_url=base_url,**kwargs)
        self.model = kwargs.get("model", "deepseek-r1:32b-qwen-distill-q8_0")
        self.base_url = base_url.rstrip('/')
        logging.info(f"OpenAI class initialized with base_url: {self.base_url}")

# --- Helper Function for Get Completion with Retry ---
def get_completion_with_retry(provider: OpenAI, prompt: str, system_message: str, **kwargs: None) -> str:
    """Wrapper function to retry get_completion in case of ReadError."""
    try:
        response_content, _ = provider.get_completion(prompt, system_message, **kwargs) #Get response
        return response_content #Return response
    except Exception as e:
        logging.error(f"Error in get_completion_with_retry: {e}")
        return ""

# --- Glossary Formatter ---
def format_glossary(glossary: List[Dict[str, str]]) -> str:
    """Format the glossary to a xml format."""
    glossary_xml = "".join(
        f"<GlossaryItem><Source>{item['source']}</Source><Target>{item['target']}</Target><Notes>{item['notes']}</Notes></GlossaryItem>"
        for item in glossary
    )
    return glossary_xml
# --- Helper Class for LLM Provider Wrapper ---
class LLMProviderWrapper(LLMProvider):
        def __init__(self, openai_instance):
            self.openai_instance = openai_instance
            self.model = self.openai_instance.model

        def num_tokens_in_string(self, input_str: str) -> int:
            return self.openai_instance.num_tokens_in_string(input_str)

        def split_text(self, text: str, max_tokens: int = 1024) -> List[str]:
            return self.openai_instance.split_text(text, max_tokens)

        def get_completion(self, prompt: str, system_message: str = "You are a helpful assistant.") -> Tuple[str, Usage]:
            return self.openai_instance.get_completion(prompt, system_message)
# --- Prompt Templates ---
PROMPT_INITIAL_TRANSLATION_TITLE_SYSTEM = (
    "You are an expert translator fluent in {source_lang} and {target_lang}. You are an expert of translating titles of philosophical and spiritual book by keeping nuances intact."
    "Translate the user's input title accurately and fluently into {target_lang}. The translation should be in PhD. level. The translation should be very short and concise. Do not add additional information or explanations."
    "Use the provided category to guide your translation, but focus on translating the title directly without adding extra context. "
    "Document category: {document_category} "
    "Output ONLY the translated title, and nothing else. Do not include any introductory or concluding remarks."
)

PROMPT_INITIAL_TRANSLATION_TITLE = """Translate the following title from {source_lang} to {target_lang}:\\\\\\\\n\\\\\\\\n{chunk_to_translate}\\\\\\\\n\\\\\\\\nOutput only the translation of the text above."""


PROMPT_QUALITY_CHECK_TITLE_SYSTEM = (
    "You are an expert in assessing the quality of {target_lang} translations for titles. Your task is to meticulously identify errors, inconsistencies, and areas for improvement in the translated title."
    "Use the provided category to guide your quality check, but ONLY focus on linguistic and translation errors."
    "Document category: {document_category} "
    "You must perform a detailed linguistic review of the translation. Check for accuracy, grammar, syntax, and cultural appropriateness."
    "Ensure that the translation follows the conventions of {target_lang} writing and avoids any unnatural or awkward phrasing. Make sure that the translation sounds natural and fluent as if the title was initially created in {target_lang}."
     "Identify areas that require improvement and provide detailed suggestions. Do not rewrite the translation, and DO NOT provide alternative translations. Do not include any explanations or interpretations of the title, provide only suggestions to improve the title."
    "DO NOT SUGGEST ADDING ADDITIONAL INFORMATION OR PARTS INTO THE TEXT."
    "Output ONLY the suggestions, and nothing else. Do not include any introductory or concluding remarks."
)

PROMPT_QUALITY_CHECK_TITLE = """Please check the quality of this {target_lang} title translation:\\\\\\\\\\\\\\\\n\\\\\\\\\\\\\\\\n{improved_text}\\\\\\\\\\\\\\\\n\\\\\\\\\\\\\\\\nand provide your feedback. Output only the suggestions."""

PROMPT_REVISE_TRANSLATION_TITLE_SYSTEM = (
   "You are an expert in revising and finalizing {target_lang} translations for titles."
   "Your task is to take the translated title and suggestions, and create a final version of the title translation that is accurate, fluent, and natural."
   "Use the provided category to guide your revision, but only focus on making the title better based on the feedback."
   "Document category: {document_category} "
   "Use the suggestions and improve the translation accordingly. Make sure to incorporate the necessary changes, corrections, and adjustments to reflect the best possible translated title."
   "Your goal is to achieve a translation that reads as if the title was initially written in {target_lang}, avoiding any unnatural or awkward phrasing. "
    "Focus on improving the consistency, the accuracy, and the overall readability of the title. "
    "Output ONLY the revised translation, do not add any additional text, label or explanations. Do not include any introductory or concluding remarks."
)

PROMPT_REVISE_TRANSLATION_TITLE = """Revise the {target_lang} title translation based on the feedback.\\\\\\\\\\\\\\\\n\\\\\\\\\\\\\\\\nTranslation: {improved_text}\\\\\\\\\\\\\\\\n\\\\\\\\\\\\\\\\nFeedback: {quality_feedback}\\\\\\\\\\\\\\\\n\\\\\\\\\\\\\\\\nOutput only the revised translation."""

PROMPT_INITIAL_TRANSLATION_SYSTEM = (
    """You are an expert translator fluent in {source_lang} and {target_lang}, specializing in philosophical and spiritual texts with a PhD-level understanding.

Approach this translation through the following continuous thought process:

1. Initial Text Analysis:
- Carefully read and absorb the source text
- Identify key philosophical or spiritual concepts
- Note any specialized terminology or nuanced expressions
- Consider the text's tone and stylistic elements

2. Context Integration:
Document Context:
- Category: {document_category}
- Summary: {document_summary}
- Reflect on how this context influences word choices and interpretation

3. Translation Development:
- Begin with a preliminary translation
- Pay special attention to:
  • Philosophical terminology accuracy
  • Spiritual concept preservation
  • Cultural nuances
  • Academic-level vocabulary
- Consider multiple alternatives for complex terms

4. Refinement Process:
- Review the translation for:
  • Conceptual accuracy
  • Philosophical depth
  • Academic rigor
  • Natural flow in {target_lang}
- Adjust any terms that could be enhanced

5. Final Polish:
- Ensure PhD-level language quality
- Verify that spiritual and philosophical nuances are preserved
- Confirm alignment with document context
- Make final adjustments for fluency"""
)

PROMPT_INITIAL_TRANSLATION = """Now, applying this thought process, translate the following text from {source_lang} to {target_lang}:\\\\\\\\n\\\\\\\\n{chunk_to_translate}\\\\\\\\n\\\\\\\\nOutput only the translation of the text above."""

PROMPT_IMPROVE_TRANSLATION_SYSTEM = (
    """You are an expert in improving the quality of {target_lang} translations, with a PhD-level understanding of linguistic nuances and translation theory.

Approach this translation review through the following continuous thought process:

1. Initial Assessment:
- Read both the original text and translation carefully
- Note your first impressions of flow and naturalness
- Identify any immediately apparent issues
- Consider the document context:
  • Category: {document_category}
  • Summary: {document_summary}

2. Linguistic Analysis:
- Examine the following elements:
  • Grammar structure and accuracy
  • Verb conjugations and tenses
  • Case endings and particles
  • Sentence structure and flow
- Compare against native {target_lang} patterns
- Note any deviations from natural language use
- Examine word choice patterns:
  • Identify words that feel translated
  • Look for more natural alternatives
  • Check for register appropriateness
  • Verify collocational accuracy
- Review expression patterns:
  • Assess phrase naturalness
  • Check for calques and literal translations
  • Identify opportunities for idiomatic expressions
  • Verify natural word order

3. Cultural and Contextual Review:
- Evaluate idiom and expression usage
- Check for literal translations that could be improved
- Consider cultural appropriateness
- Verify consistency with document context
- Look for opportunities to use native {target_lang} expressions

4. Precision Assessment:
- Review technical terminology
- Check for accuracy in specialized vocabulary
- Verify meaning preservation
- Ensure no information is added or removed
- Analyze flow and rhythm:
  • Check sentence patterns
  • Verify natural transitions
  • Assess paragraph coherence
  • Review overall text flow
- Examine clarity and precision:
  • Identify any awkward phrasing
  • Check for overcomplicated structures
  • Verify information hierarchy
  • Assess logical connections

5. Improvement Formulation:
- Focus only on necessary changes
- Prioritize suggestions by importance
- Keep recommendations specific and actionable
- Maintain original text structure
- Avoid suggesting additional content

Based on this analysis, provide your suggestions in the following format:

1. Word Choice and Expressions:
[List specific suggestions for more natural vocabulary and phrases]

2. Cultural and Register Adjustments:
[List recommendations for cultural adaptation and register refinement]

3. Flow and Readability:
[List improvements for better text flow and natural rhythm]

Remember:
- Focus on making the text sound natural in {target_lang}
- Maintain the original meaning while adapting expressions
- Consider target audience and cultural context
- Preserve technical accuracy while improving flow
- Focus only on improving existing content
- Do not suggest adding new information
- Keep suggestions precise and specific
- Maintain the original meaning and scope"""
)

PROMPT_IMPROVE_TRANSLATION = """Now, analyze the following translation and provide suggestions for improving this {target_lang} translation:\\\\\\\\\\\\\\\\n\\\\\\\\\\\\\\\\n{translated_text}\\\\\\\\\\\\\\\\n\\\\\\\\\\\\\\\\nOutput only the suggestions."""

PROMPT_QUALITY_CHECK_SYSTEM = (
    """You are an expert in assessing the quality of {target_lang} translations, with extensive experience in linguistic analysis and quality assurance.

Approach this quality assessment through the following continuous thought process:

1. Context Understanding:
- Review the document context:
  • Category: {document_category}
  • Summary: {document_summary}
- Consider the expected style and tone
- Note any specialized terminology requirements
- Identify key concepts that must be preserved

2. Natural Language Assessment:
- Evaluate word choice naturalness:
  • Check if words feel forced or overly formal
  • Identify any direct translations that sound unnatural
  • Look for more natural alternatives to technical terms
  • Assess if vocabulary matches native speaker intuition
- Examine collocation patterns:
  • Verify natural word combinations
  • Check if verb-noun pairs sound natural
  • Assess adjective-noun combinations
  • Review adverb placement and usage
- Analyze expression patterns:
  • Identify any non-native phrasing
  • Check for calques (word-for-word translations)
  • Look for more natural equivalent expressions
  • Verify idiomatic usage

3. Register and Style Consistency:
- Evaluate language register appropriateness:
  • Check if formality level matches context
  • Verify consistency in honorific usage
  • Assess politeness level appropriateness
  • Review specialized terminology usage
- Examine stylistic elements:
  • Check for natural rhythm and flow
  • Verify sentence length variation
  • Assess paragraph transitions
  • Review overall text cohesion

4. Accuracy Analysis:
- Compare source and target texts closely
- Check for:
  • Meaning preservation
  • Concept accuracy
  • Technical terminology
  • Numerical consistency
- Note any meaning distortions or omissions

5. Linguistic Assessment:
- Evaluate grammar elements:
  • Verb forms and tenses
  • Particle usage
  • Sentence structure
  • Punctuation
- Examine syntax patterns
- Review collocation usage
- Check for register consistency

6. Cultural and Natural Flow Review:
- Assess cultural appropriateness
- Evaluate naturalness in {target_lang}
- Check for:
  • Idiomatic expressions
  • Cultural references
  • Writing conventions
- Identify any foreign-sounding phrases

7. Professional Standards Check:
- Verify adherence to:
  • Target language conventions
  • Industry-specific terminology
  • Style consistency
  • Formatting standards

Quality Report Formation:
Structure your findings as follows:

A. Word Choice and Collocations:
[List issues with unnatural word combinations and suggestions for more native-like expressions]

B. Expression Patterns:
[List instances of non-native phrasing and recommended natural alternatives]

C. Register and Style:
[List any inconsistencies in register and style with suggested improvements]

D. Critical Issues:
[List issues that significantly impact meaning or understanding]

E. Language Accuracy:
[List grammar, syntax, and structural issues]

F. Natural Flow:
[List areas where the text could sound more natural]

G. Style and Consistency:
[List style-related improvements needed]

Remember:
- Focus on identifying existing issues
- Do not suggest adding new content
- Provide specific, actionable feedback
- Maintain focus on the original scope"""
)

PROMPT_QUALITY_CHECK = """Now, assess the following {target_lang} translation:\\\\\\\\\\\\\\\\n\\\\\\\\\\\\\\\\n{improved_text}\\\\\\\\\\\\\\\\n\\\\\\\\\\\\\\\\nand provide your feedback. Output only the suggestions."""

PROMPT_REVISE_TRANSLATION_SYSTEM = (
   """You are an expert in revising and finalizing {target_lang} translations, with deep expertise in creating publication-ready content.

Approach this revision through the following continuous thought process:

1. Holistic Review:
- Examine the original translation and improvement suggestions
- Consider the document context:
  • Category: {document_category}
  • Summary: {document_summary}
- Identify patterns in suggested improvements
- Note key areas requiring attention

2. Systematic Improvement:
- Address improvements in order of importance:
  • Critical meaning corrections
  • Grammar and syntax refinements
  • Style and flow enhancements
  • Cultural adaptations
- Consider how each change affects the whole text
- Maintain consistency throughout revisions

3. Natural Language Enhancement:
- Refine the text to sound native to {target_lang}
- Consider:
  • Natural word order
  • Appropriate idioms
  • Cultural references
  • Local expressions
- Ensure consistent tone and style

4. Technical Precision:
- Verify specialized terminology
- Check for:
  • Technical accuracy
  • Field-specific conventions
  • Consistent terminology
  • Proper formatting

5. Final Polish:
- Read the text as if it were originally written in {target_lang}
- Focus on:
  • Overall flow
  • Natural transitions
  • Consistent voice
  • Professional tone

6. Quality Assurance:
- Verify that all suggestions have been appropriately addressed
- Ensure no new issues were introduced
- Confirm adherence to target language conventions
- Check for consistency in revisions

Remember:
- Focus on creating a polished final version
- Maintain the original meaning and scope
- Ensure natural flow in {target_lang}
- Preserve technical accuracy
- Output ONLY the revised translation 
- Do not add any additional text, label or explanations 
- Do not include any introductory or concluding remarks."""
)

PROMPT_REVISE_TRANSLATION = """Now, revise the following {target_lang} translation, considering the provided suggestions.\\\\\\\\\\\\\\\\n\\\\\\\\\\\\\\\\nTranslation: {improved_text}\\\\\\\\\\\\\\\\n\\\\\\\\\\\\\\\\nSuggestions: {quality_feedback}\\\\\\\\\\\\\\\\n\\\\\\\\\\\\\\\\nOutput only the revised translation."""
#############################################################
PROMPT_SUMMARIZE_DOCUMENT_SYSTEM = (
    "You are an expert summarizer tasked with creating a detailed JSON summary of long documents. "
    "The summary MUST contain at least {max_tokens} words.\n\n"
    
    "Follow this mandatory process:\n\n"
    
    "1. Document Analysis:\n"
    "   - Break the document into major sections\n"
    "   - For each section, identify:\n"
    "     * Main arguments\n"
    "     * Key findings\n"
    "     * Important details\n"
    "     * Supporting evidence\n\n"
    
    "2. Length Enforcement:\n"
    "   - The summary MUST contain at least {max_tokens} words.\n"
    "   - Include detailed coverage of ALL major sections\n"
    "   - If your summary seems short, expand it by:\n"
    "     * Adding more supporting details\n"
    "     * Elaborating on key concepts\n"
    "     * Including relevant examples\n"
    "     * Explaining relationships between ideas\n\n"
    
    "3. Content Requirements:\n"
    "   - Begin with an overview paragraph\n"
    "   - Cover each major section in detail\n"
    "   - Include transitions between topics\n"
    "   - End with key conclusions\n"
    "   - Maintain cause-and-effect relationships\n"
    "   - Preserve critical numerical data\n"
    "   - Include methodology if present\n\n"
    
    "4. Keyword Extraction:\n"
    "   - Extract at least 15 keywords\n"
    "   - Include:\n"
    "     * Major themes\n"
    "     * Key technical terms\n"
    "     * Important concepts\n"
    "     * Significant findings\n\n"
    
    "5. Quality Check:\n"
    "   - Verify summary length meets minimum {max_tokens} words\n"
    "   - Ensure all major sections are covered\n"
    "   - Confirm key arguments are preserved\n"
    "   - Check for completeness and accuracy\n\n"
    
    "Output Format:\n"
    "{\"summary\": \"YOUR EXTENSIVE SUMMARY HERE - MUST BE AT LEAST '{max_tokens}' TOKENS\", "
    "\"keywords\": [\"keyword1\", \"keyword2\", ... ]}\n\n"
    
    "CRITICAL: If your summary is shorter than {max_tokens} words, you MUST expand it before outputting. "
    "A summary shorter than {max_tokens} words is UNACCEPTABLE. Your summary should be "
    "detailed enough that someone could understand all major points without reading the original.\n\n"
    
    "Remember: Output only valid JSON with no additional text or explanations."
)
PROMPT_SUMMARIZE_DOCUMENT = "Summarize the following text: \\\\\\\\\\\\\\\\n\\\\\\\\\\\\\\\\n{text}"
PROMPT_EXTRACT_CATEGORY_SYSTEM = "You are an expert document classifier. Your task is to classify the text based on the main topics. The category should be very specific, for example, a philosophical book investigating creation of consciousness or a sci-fi novel about transhumanism or a scientific article on LLMs and so on. Do not create new categories. Use only existing categories, You must choose a category that is most relevant to the text."
PROMPT_EXTRACT_CATEGORY = "Classify the following text: \\\\\\\\n\\\\\\\\n{text}"

def validate_llm_response(response: str) -> tuple[bool, str]:
    """
    Validates LLM response and attempts to fix common issues.
    If plain text is received, attempts to convert it to proper JSON format.
    Returns (success_boolean, json_string)
    """
    response = str(response).strip()
    logging.debug(f"Validating response: {response}")

    # Remove any markdown code block markers
    response = re.sub(r'```(?:json)?\s*', '', response, flags=re.IGNORECASE)
    response = re.sub(r'`\s*', '', response, flags=re.IGNORECASE)

    # First try to find and parse existing JSON
    potential_json = re.search(r'\{.*\}', response, re.DOTALL)
    if potential_json:
        logging.debug("FOUND JSON!!!")
        json_str = potential_json.group(0)

        # Fix common JSON formatting issues
        json_str = json_str.replace('\\n', ' ').replace('\\', '')

        # Ensure proper quote usage
        json_str = re.sub(r'(?<=\{|\,)\s*"?(\w+)"?\s*:', r'"\1":', json_str)


        try:
            json_obj = json.loads(json_str)
            logging.debug(f"Valid JSON after corrections: {json_str}")
            return True, json.dumps(json_obj, ensure_ascii=False)
        except json.JSONDecodeError as e:
            logging.debug(f"JSONDecodeError after corrections: {e}")
            pass  # Fall through to plain text handling

    # If we reach here, either no JSON was found or it was invalid
    # Try to parse as plain text and convert to JSON
    try:
        # Try to separate summary and keywords if they exist
        # First look for explicit "Keywords:" section
        keywords_match = re.split(r'\n*Keywords?:?\s*', response, flags=re.IGNORECASE)
        if len(keywords_match) > 1:
            # We found a keywords section
            summary = keywords_match[0].strip()
            keywords_text = keywords_match[1].strip()

            # Process keywords
            keywords = [k.strip() for k in re.split(r'[,;\\n]', keywords_text) if k.strip()]
        else:
            # No explicit keywords section found
            # Try to identify if there's a list-like structure at the end
            list_pattern = r'(?:[-\*•]|\d+\.)\s*\w+'
            if re.search(list_pattern, response, re.MULTILINE):
                # Split on the first list item
                parts = re.split(fr'(?m)^{list_pattern}', response, maxsplit=1)
                if len(parts) > 1:
                    summary = parts[0].strip()
                    list_items = re.findall(fr'{list_pattern}.*?(?=(?:[-\*•]|\d+\.)|$)', response, re.MULTILINE)
                    keywords = [re.sub(r'^(?:[-\*•]|\d+\.)\s*', '', item).strip() for item in list_items]
                else:
                    summary = response
                    keywords = []
            else:
                # No clear keyword section found
                summary = response
                keywords = []

        # Clean up summary
        summary = re.sub(r'^Summary:?\s*', '', summary, flags=re.IGNORECASE).strip()

        # Construct proper JSON
        json_obj = {
            "summary": summary,
            "keywords": keywords
        }
        json_str = json.dumps(json_obj, ensure_ascii=False)
        logging.debug(f"Converted plain text to JSON: {json_str}")
        return True, json_str

    except Exception as e:
        logging.error(f"Error converting plain text to JSON: {e}")
        # Last resort: wrap the entire text as a summary with no keywords
        try:
            json_obj = {
                "summary": response.strip(),
                "keywords": []
            }
            json_str =  json.dumps(json_obj, ensure_ascii=False)
            logging.debug(f"Last resort: wrapped text into JSON: {json_str}")
            return True, json_str
        except:
            logging.error(f"Error wrapping text into json {e}")
            return False, response


def summarize_document(provider: LLMProviderWrapper, text_chunks: list[str], max_tokens_1: int = 250, max_tokens_2: int = 1250, max_keywords: int = 50) -> dict:
    """Enhanced document summarization with better error handling and format enforcement."""
    logging.info("Starting document summarization...")
    final_summary = ""
    all_keywords = []
    temp_dir = "temp_summarization_files"
    os.makedirs(temp_dir, exist_ok=True)
    unique_id = str(uuid.uuid4())[:8]
    max_tokens=500

    # First pass: Get summaries of chunks
    for i, chunk in enumerate(tqdm(text_chunks, desc="Summarizing chunks")):
        logging.info(f"Processing chunk {i+1}/{len(text_chunks)}: {chunk[:100]}...")
        messages = [{
            "content": f"Summarize this text as a JSON object: {chunk}. max_tokens: {max_tokens_1}",
            "system_message": PROMPT_SUMMARIZE_DOCUMENT_SYSTEM
        }]

        e = None
        try:
            response = _call_llm(provider, messages)
            if isinstance(response, dict) and "content" in response:
                try:
                    summary_text = response["content"]
                    if not summary_text:
                        logging.warning(f"Chunk {i+1}: Empty response content received.")
                        continue
                    logging.info("Starting validate llm...")
                    is_valid, json_str = validate_llm_response(summary_text)
                    if is_valid:
                        logging.info(f"Chunk {i+1}: Valid response!")
                        try:
                            summary_json = json.loads(json_str)
                            logging.info(
                                f"Chunk {i+1} summary snippet: {summary_json.get('summary', '')[:100]}")
                            final_summary += summary_json.get("summary", "") + " "
                            all_keywords.extend(summary_json.get("keywords", []))
                        except json.JSONDecodeError as err:
                            e = err
                            logging.error(f"Chunk {i+1}: Failed to parse validated JSON: {e}")
                    else:
                        logging.warning(f"Chunk {i+1}: Invalid json format, skipped.")
                except Exception as err:
                    e = err
                    logging.error(f"Error processing chunk {i+1} first if: {e}")
            else:
                try:
                    logging.error(f"Chunk {i+1}: Did not receive a valid dictionary from LLM")
                except Exception as err:
                     e = err
                     logging.error(f"Error processing chunk {i+1} else condition: {e}")
        except Exception as err:
             e = err
             logging.error(f"Error processing chunk {i+1}: {e}")
        if e is not None:
            continue

    # Process accumulated keywords before second pass
    keyword_counts = Counter(all_keywords)
    top_keywords = [kw for kw, _ in keyword_counts.most_common(max_keywords)]

    # Save the final_summary to a temp file
    final_summary_file = os.path.join(temp_dir, f"final_summary_{unique_id}.txt")
    try:
        with open(final_summary_file, "w", encoding="utf-8") as f:
            f.write(final_summary)
        logging.debug(f"Saved final_summary to {final_summary_file}")
    except Exception as e:
        logging.error(f"Error saving final summary file: {e}")

    # Second pass: Final summary with explicit minimum word count
    if final_summary:
        messages = [{
            "content": (
                f"Create a detailed final JSON summary of this text that MUST be at least 250 words long: {final_summary}. "
                "Ensure you include extensive detail and context for all major points. "
                "The summary should be comprehensive enough that someone could understand ALL major points "
                "without reading the original text. If needed, elaborate on implications and relationships "
                "between different aspects of the content. min_tokens: {max_tokens_2}\n"
                "Remember: Output ONLY valid JSON in the format "
                '{"summary": "text", "keywords": ["word1", "word2"]}'
            ),
            "system_message": PROMPT_SUMMARIZE_DOCUMENT_SYSTEM
        }]

        try:
            response = _call_llm(provider, messages)
            if isinstance(response, dict) and "content" in response:
                # Save the LLM's response before parsing
                llm_response_file = os.path.join(temp_dir, f"llm_response_{unique_id}.txt")
                try:
                    with open(llm_response_file, "w", encoding="utf-8") as f:
                        f.write(response["content"])
                    logging.debug(f"Saved LLM response to {llm_response_file}")
                except Exception as e:
                    logging.error(f"Error saving LLM response file: {e}")

                is_valid, json_str = validate_llm_response(response["content"])
                if is_valid:
                    logging.info("VALID RESPONSE!!!!!!...")
                    try:
                        second_pass_json = json.loads(json_str)
                        summary = second_pass_json.get("summary", "").strip()
                        
                        # Check if summary meets minimum word count
                        word_count = len(summary.split())
                        logging.info(f"Summary word count: {word_count}")
                        
                        if word_count < 250:
                            logging.warning(f"Summary too short ({word_count} words). Requesting expansion...")
                            # Request expansion if too short
                            expansion_messages = [{
                                "content": (
                                    f"Expand this summary to be at least 500 words while maintaining accuracy and adding detail: {summary}"
                                ),
                                "system_message": PROMPT_SUMMARIZE_DOCUMENT_SYSTEM
                            }]
                            expansion_response = _call_llm(self.llm_provider, expansion_messages)
                            if isinstance(expansion_response, dict) and "content" in expansion_response:
                                
                                llm_response_file_ex = os.path.join(temp_dir, f"llm_response_ex{unique_id}.txt")
                                try:
                                    with open(llm_response_file_ex, "w", encoding="utf-8") as f:
                                        f.write(expansion_response["content"])
                                    logging.debug(f"Saved LLM response to {llm_response_file_ex}")
                                except Exception as e:
                                    logging.error(f"Error saving LLM response file expanded: {e}")
                                
                                is_valid, expanded_json_str = validate_llm_response(expansion_response["content"])
                                if is_valid:
                                    expanded_json = json.loads(expanded_json_str)
                                    summary = expanded_json.get("summary", summary)
                        
                        return {
                            "summary": summary,
                            "keywords": top_keywords
                        }
                    except json.JSONDecodeError as e:
                        logging.error(f"Failed to parse second pass JSON {e}")
        except Exception as e:
            logging.error(f"Error in final summarization: {e}")

    # Fallback response
    return {
        "summary": final_summary.strip(),
        "keywords": top_keywords
    }

def extract_category(provider: LLMProviderWrapper, text_chunks: List[str]) -> str:
    """Extracts the category of the input text using the LLM on chunks and then combine"""
    logging.info("Starting document category extraction...")
    all_categories = []
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1024, chunk_overlap=0, separators=["\n\n", "\n", "."])

    for chunk in tqdm(text_chunks, desc="Categorizing chunks"):
        messages = [{"content": PROMPT_EXTRACT_CATEGORY.format(text=chunk), "system_message": PROMPT_EXTRACT_CATEGORY_SYSTEM}]
        try:
            response = _call_llm(provider, messages)
            if isinstance(response, dict) and "content" in response:
                category = response["content"].strip()
                all_categories.append(category)

            else:
                logging.error("Error: Did not get proper response from LLM for category extraction")

        except Exception as e:
            logging.error(f"Error during chunk category extraction: {e}")
    # Now do the final category extraction
    all_categories_str = ", ".join(all_categories)
    messages = [{"content": PROMPT_EXTRACT_CATEGORY.format(text=all_categories_str), "system_message": PROMPT_EXTRACT_CATEGORY_SYSTEM}]
    try:
        response = _call_llm(provider, messages)
        if isinstance(response, dict) and "content" in response:
            category = response["content"].strip()
            logging.info(f"Document category extraction complete. Category: {category}")
            return category
        else:
            logging.error("Error: Did not get proper response from LLM for category extraction")
            return "Uncategorized"
    except Exception as e:
        logging.error(f"Error during final category extraction: {e}")
        return "Uncategorized"
##########################################################################

openai_api_key = 'ollama'
base_url = 'https://b125-34-143-154-113.ngrok-free.app'
# Initialize the OpenAI class here
provider = OpenAI(api_key=openai_api_key, model="deepseek-r1:32b-qwen-distill-q8_0", base_url=base_url)
# --- Autogen Agent Configuration ---
# Update the config_list to disable OpenAI-style endpoints
config_list = [
    {
        "model": "deepseek-r1:32b-qwen-distill-q8_0",
        "base_url": "https://b125-34-143-154-113.ngrok-free.app/v1",
        "api_key": "ollama",
    }
]

# Enhanced ollama_generate function with debug logging
def _call_llm(provider: LLMProviderWrapper, messages: List[Dict], max_tokens: int = 100000) -> Dict:
    """Helper function to call the LLM and handle errors."""
    message = messages[0]["content"] if messages and "content" in messages[0] else ""
    system_message = messages[0].get("system_message", "")
    try:
        response_content, usage = provider.get_completion(message, system_message)
        if not response_content:
            logging.error("Empty response content received from provider.get_completion")
            raise ValueError("Empty response content received from provider.get_completion")
        logging.info(f"Successfully generated response: {response_content[:100]}...")
        if not isinstance(response_content, str):
            response_content = str(response_content) # Ensure response content is a string

        return {"content": response_content, "usage": usage.to_dict()}

    except Exception as e:
        logging.error(f"Error in _call_llm: {str(e)}", exc_info=True)
        raise

def ollama_generate(messages, **kwargs):
    """Enhanced ollama_generate function with better error handling and logging"""
    logging.info(f"ollama_generate called with messages: {messages}")
    try:
        openai_api_key = 'ollama'
        base_url = 'https://b125-34-143-154-113.ngrok-free.app'  # Update with your current URL
        provider = OpenAI(api_key=openai_api_key, model="deepseek-r1:32b-qwen-distill-q8_0", base_url=base_url)
        return _call_llm(provider, messages, **kwargs)
    except Exception as e:
         logging.error(f"Error in ollama_generate: {str(e)}", exc_info=True)
         raise

def _agent_call_with_retry(self, user_proxy, agent, messages, log_message, max_retries=3, max_consecutive_empty=2):
        """Helper function to call an agent with retry logic and cycle detection."""
        empty_responses = 0

        for attempt in range(max_retries):
            logging.info(f"{log_message}, attempt {attempt + 1}/{max_retries}")
            logging.info(f"Initiating chat with message: {messages[0]['content'][:100]}...")

            try:

                logging.info(f"Calling ollama_generate with messages: {messages}")

                # Call ollama_generate with explicit parameters
                response = _call_llm(agent.llm_config['config_list'][0]['provider'](api_key=agent.llm_config['config_list'][0]['api_key'], base_url=agent.llm_config['config_list'][0]['base_url']), messages, max_tokens=100000)

                logging.info(f"Raw response from ollama_generate: {response}")

                if response is None:
                    logging.error("Received None response from ollama_generate")
                    continue

                if isinstance(response, dict):
                    content = response.get("content", "")
                    if content and len(content.strip()) > 0:
                        logging.info(f"{log_message} successful. Response: {content[:100]}...")
                        return content
                    else:
                        logging.warning("Response dictionary has no content or empty content")
                else:
                    logging.error(f"Unexpected response type: {type(response)}")

                empty_responses += 1
                if empty_responses >= max_consecutive_empty:
                    logging.error(f"Too many consecutive empty responses ({empty_responses}). Breaking retry loop.")
                    return f"Error: Multiple empty responses for {log_message}"

                logging.warning(f"{log_message} returned empty response. Retrying...")
                time.sleep(2)  # Increased sleep time between retries

            except Exception as e:
                logging.error(f"Error during chat: {str(e)}", exc_info=True)  # Added full exception info
                time.sleep(2)
                continue

        logging.error(f"{log_message} failed after {max_retries} attempts.")
        return f"Error: Failed to get response after {max_retries} attempts for {log_message}"

# Modified llm_config
llm_config_initial = {
    "timeout": 600,
    "cache_seed": 42,
    "temperature": 0.3,
    "frequency_penalty": 0.2,
    "presence_penalty": 0.1,
    "top_k": 40,
    "config_list": config_list,
}

llm_config_quality = {
    "timeout": 600,
    "cache_seed": 42,
    "temperature": 0.2,
    "top_p": 0.8,
    "frequency_penalty": 0.1,
    "presence_penalty": 0.1,
    "top_k": 30,
    "config_list": config_list,
}

llm_config_revision = {
    "timeout": 600,
    "cache_seed": 42,
    "temperature": 0.4,
    "top_p": 0.92,
    "frequency_penalty": 0.3,
    "presence_penalty": 0.2,
    "top_k": 45,
    "config_list": config_list,
}

llm_config_improver = {
    "timeout": 600,
    "cache_seed": 42,
    "temperature": 0.5,
    "top_p": 0.95,
    "frequency_penalty": 0.4,
    "presence_penalty": 0.3,
    "top_k": 50,
    "config_list": config_list,
}

# Updated create_agent function with debug logging
def create_agent(name, system_message="None", llm_config=None, functions=None):
    logging.info(f"Creating agent {name} with system message: {system_message[:100]}...")
    agent = autogen.AssistantAgent(
        name=name,
        system_message=system_message,
        llm_config=llm_config,
    )
    logging.info(f"Agent {name} created with config: {agent.llm_config}")
    return agent

def create_user_proxy(name, human_input_mode="NEVER"):
    return autogen.UserProxyAgent(
        name=name,
        human_input_mode=human_input_mode,
        max_consecutive_auto_reply=0
    )

def recursive_structure_aware_chunking(text: str, max_tokens: int = 1024) -> List[str]:
    """Splits text into chunks using a recursive structure-aware approach."""
    # First, attempt to partition using unstructured to get elements
    # Here, we're skipping unstructured because it is causing issue with text input
    # Fallback to recursive character splitting if partitioning fails
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=max_tokens, chunk_overlap=0, separators=["\n\n", "\n", "."])
    return text_splitter.split_text(text)


class SentimentChunker:
    def __init__(self, max_tokens: int = 1024):
        self.max_tokens = max_tokens

    def create_chunks(self, text: str) -> List[Tuple[str, int]]:
        """Create chunks based on the recursive structure-aware method."""
        chunks = recursive_structure_aware_chunking(text, self.max_tokens)
        return [(chunk, i) for i, chunk in enumerate(chunks)]

class UniversalTranslator:
    def __init__(
            self,
            source_language: str,
            target_language: str,
            llm_provider: LLMProvider,
            output_file: str,
            agent_pipeline: List[str] = ["translator", "improver", "quality_checker", "reviser"],
             title_agent_pipeline: List[str] = ["initial_translator_title", "quality_checker_title", "reviser_title"],
            **kwargs,
    ):
        self.source_language = source_language
        self.target_language = target_language
        self.llm_provider = llm_provider
        self.output_file = output_file
        self.chunks_dict = OrderedDict()  # Store chunks with their order
        self.chunker = SentimentChunker(max_tokens=kwargs.get("max_tokens", 1024))
        self.max_concurrent_requests = kwargs.get("max_concurrent_requests", 5)
        self.agent_pipeline = agent_pipeline #Store the pipeline
        self.title_agent_pipeline = title_agent_pipeline
        self.document_summary = None  # Initialize document summary
        self.document_category = None  # Initialize document category
        self.translator_agent = create_agent(
            name="Translator",
            llm_config=llm_config_initial,
        )
        self.improver_agent = create_agent(
            name="Translation_Improver",
            llm_config=llm_config_improver,
        )

        self.quality_checker_agent = create_agent(
            name="Quality_Checker",
            llm_config=llm_config_quality,
        )

        self.reviser_agent = create_agent(
            name="Reviser",
            llm_config=llm_config_revision,
        )
        
        self.initial_translator_title_agent = create_agent(
            name="Initial_Translator_Title",
            llm_config=llm_config_initial,
        )
        self.quality_checker_title_agent = create_agent(
            name="Quality_Checker_Title",
            llm_config=llm_config_quality,
        )

        self.reviser_title_agent = create_agent(
            name="Reviser_Title",
            llm_config=llm_config_revision,
        )
        self.min_tokens_for_full_pipeline = kwargs.get("min_tokens_for_full_pipeline", 100) #Set the min tokens
        self.title_post_process = kwargs.get("title_post_process", None)
        #self.glossary = {} #Initialize the glossary

    def translate(self, file_path: str, post_processing: List[Callable] = []) -> List[Dict]:
            """Translates the text using structure-aware chunking and ordered processing."""
            logging.info("Starting structure-aware translation...")

            # Read the structured text from the file
            structured_text = read_file(file_path)
            if not structured_text:
                logging.error("No text provided to translate")
                return []

            # Combine all the contents into one big string for summarization and category
            combined_text = ""
            for element in structured_text:
                combined_text += element["content"] + " "

            # Split text into chunks before summarization and categorization
            text_splitter = RecursiveCharacterTextSplitter(chunk_size=1024, chunk_overlap=0, separators=["\n\n", "\n", "."]) #Or whatever chunk size you want
            text_chunks = text_splitter.split_text(combined_text)


            # Summarize and categorize the document
            document_summary = summarize_document(self.llm_provider, text_chunks)
            document_category = extract_category(self.llm_provider, text_chunks)
            logging.info(f"Document Summary: {document_summary}")
            logging.info(f"Document Category: {document_category}")

            # Set the document summary and category to the instance variables
            self.document_summary = document_summary
            self.document_category = document_category

            # Concatenate consecutive paragraphs
            concatenated_elements = []
            current_paragraph = ""
            for element in structured_text:
                if element.get("type") == "Paragraph":
                    current_paragraph += element["content"] + " "
                else:
                    if current_paragraph:
                        concatenated_elements.append({"type": "Paragraph", "content": current_paragraph.strip()})
                        current_paragraph = ""
                    concatenated_elements.append(element) # Append other elements directly

            #Append any remaining paragraphs
            if current_paragraph:
                concatenated_elements.append({"type": "Paragraph", "content": current_paragraph.strip()})

            # Create Glossary for the whole document
            logging.info("Generating full document glossary...")

            # Process concatenated elements and use recursive chunking where it is needed
            translated_elements = []
            total_chunks = len(concatenated_elements)
            for index, element in enumerate(concatenated_elements):
                logging.info(f"Processing element: {element.get('type', 'Unknown')}")
                try:
                    if element.get("type") == "Paragraph":
                        #Apply recursive chunking to the concatenated paragraphs
                        chunks = recursive_structure_aware_chunking(element["content"], max_tokens=1024) #Or whatever you want
                        for i, chunk in enumerate(chunks):
                            translated_element = element.copy()
                            translated_element["content"] = chunk
                            final_translated_text = self._process_translation_pipeline(chunk, translated_element, index, total_chunks)
                            translated_element["content"] = final_translated_text
                            translated_elements.append(translated_element)

                    else:
                        translated_element = element.copy()
                        final_translated_text = self._process_translation_pipeline(element["content"], element, index, total_chunks)
                        translated_element["content"] = final_translated_text
                        translated_elements.append(translated_element)


                except Exception as e:
                    error_message = f"Error processing element: {element.get('type', 'Unknown')}: {e}"
                    logging.error(error_message)
                    translated_elements.append({"type": "Error", "content": f"Error in element {element.get('type', 'Unknown')}: {e}"})



            # Apply post-processing
            final_text =  translated_elements
            for post_process in post_processing:
                final_text = post_process(final_text)

            return final_text

    def _process_translation_pipeline(self, translated_text: str, element: Dict, chunk_index: int, total_chunks: int) -> str: #, glossary: dict
            """Process a single chunk through the improvement pipeline while maintaining order."""

            improver_feedback = ""
            quality_feedback = ""
            # Check if the current element is a title or if it has less than min tokens for full pipeline
            if element.get("type", "Other") == "Title":
            #Apply only the initial translation
                # Process the title through the title pipeline
                if self.title_agent_pipeline:
                    for agent_name in self.title_agent_pipeline:
                        if agent_name == "initial_translator_title":
                            messages = [{
                                    "content": PROMPT_INITIAL_TRANSLATION_TITLE.format(
                                                            source_lang=self.source_language,
                                                            target_lang=self.target_language,
                                                            chunk_to_translate=translated_text,
                                                        ),
                                    "system_message": PROMPT_INITIAL_TRANSLATION_TITLE_SYSTEM.format(
                                        source_lang=self.source_language,
                                        target_lang=self.target_language,
                                        #document_summary = self.document_summary.get("summary", ""),
                                        document_category = self.document_category
                                        )
                                }]
                            response = _call_llm(self.llm_provider, messages)
                            if isinstance(response, dict) and "content" in response:
                                translated_text = response["content"]
                                logging.info(f"Current text after translator (title): {translated_text[:100]}...")
                            else:
                                translated_text = ""

                # Process the title through the title pipeline
                #if self.title_agent_pipeline:
                #    for agent_name in self.title_agent_pipeline:
                        elif agent_name == "quality_checker_title":
                            messages = [{
                                    "content": PROMPT_QUALITY_CHECK_TITLE.format(
                                        target_lang=self.target_language, 
                                        improved_text=translated_text
                                        ),
                                    "system_message": PROMPT_QUALITY_CHECK_TITLE_SYSTEM.format(
                                        target_lang=self.target_language,
                                        #document_summary = self.document_summary.get("summary", ""),
                                        document_category = self.document_category
                                        )
                                }]
                            response = _call_llm(self.llm_provider, messages)
                            if isinstance(response, dict) and "content" in response:
                                quality_feedback = response["content"]
                                logging.info(f"Current text after title quality checker: {quality_feedback[:100]}...")

                            else:
                                quality_feedback = ""
                        elif agent_name == "reviser_title":
                            messages = [{
                                    "content": PROMPT_REVISE_TRANSLATION_TITLE.format(
                                        target_lang=self.target_language,
                                        improved_text=translated_text,
                                        quality_feedback=quality_feedback
                                    ),
                                    "system_message": PROMPT_REVISE_TRANSLATION_TITLE_SYSTEM.format(
                                        target_lang=self.target_language,
                                        #document_summary = self.document_summary.get("summary", ""),
                                        document_category = self.document_category
                                    )
                                }]
                            response = _call_llm(self.llm_provider, messages)
                            if isinstance(response, dict) and "content" in response:
                                translated_text = response["content"]
                                logging.info(f"Current text after title reviser: {translated_text[:100]}...")
                            else:
                                translated_text = ""

                    if self.title_post_process:
                        translated_text = self.title_post_process(translated_text)
                    return translated_text # Return after initial translation

            for agent_name in self.agent_pipeline:
                if agent_name == "translator":
                    nontranslated=translated_text
                    messages = [{
                        "content": PROMPT_INITIAL_TRANSLATION.format(
                                                source_lang=self.source_language,
                                                target_lang=self.target_language,
                                                chunk_to_translate=translated_text,
                                            ),
                        "system_message": PROMPT_INITIAL_TRANSLATION_SYSTEM.format(
                            source_lang=self.source_language,
                            target_lang=self.target_language,
                            document_summary = self.document_summary.get("summary", ""),
                            document_category = self.document_category
                            )
                    }]
                    response = _call_llm(self.llm_provider, messages)

                    if isinstance(response, dict) and "content" in response:
                        translated_text = response["content"]
                        logging.info(f"Current text after translator: {translated_text[:100]}...")
                    else:
                        translated_text = ""
                elif agent_name == "improver":
                    messages = [{
                            "content": PROMPT_IMPROVE_TRANSLATION.format(
                                target_lang=self.target_language,
                                translated_text=translated_text
                                ),
                            "system_message": PROMPT_IMPROVE_TRANSLATION_SYSTEM.format(
                                target_lang=self.target_language,
                                document_summary = self.document_summary.get("summary", ""),
                                document_category = self.document_category
                                )
                        }]

                    response = _call_llm(self.llm_provider, messages)
                    if isinstance(response, dict) and "content" in response:
                        improver_feedback = response["content"]
                        logging.info(f"Current text after improver: {improver_feedback[:100]}...")
                    else:
                        improver_feedback = ""
                elif agent_name == "reviser_improver":
                    messages = [{
                            "content": PROMPT_REVISE_TRANSLATION.format(
                                target_lang=self.target_language,
                                improved_text=translated_text,
                                quality_feedback=improver_feedback
                            ),
                            "system_message": PROMPT_REVISE_TRANSLATION_SYSTEM.format(
                                target_lang=self.target_language,
                                document_summary = self.document_summary.get("summary", ""),
                                document_category = self.document_category
                            )
                            }]
                    response = _call_llm(self.llm_provider, messages)
                    if isinstance(response, dict) and "content" in response:
                        translated_text = response["content"]
                        logging.info(f"Current text after reviser (improver): {translated_text[:100]}...")
                    else:
                        translated_text = ""
                elif agent_name == "quality_checker":
                    messages = [{
                        "content": PROMPT_QUALITY_CHECK.format(
                            target_lang=self.target_language,
                            improved_text=translated_text
                            ),
                        "system_message": PROMPT_QUALITY_CHECK_SYSTEM.format(
                            target_lang=self.target_language,
                            document_summary = self.document_summary.get("summary", ""),
                            document_category = self.document_category
                            )
                    }]
                    response = _call_llm(self.llm_provider, messages)
                    if isinstance(response, dict) and "content" in response:
                        quality_feedback = response["content"]
                        logging.info(f"Current text after quality checker: {quality_feedback[:100]}...")
                    else:
                        quality_feedback = ""
                elif agent_name == "reviser_quality":
                    messages = [{
                            "content": PROMPT_REVISE_TRANSLATION.format(
                                target_lang=self.target_language,
                                improved_text=translated_text,
                                quality_feedback=quality_feedback
                                ),
                            "system_message": PROMPT_REVISE_TRANSLATION_SYSTEM.format(
                                target_lang=self.target_language,
                                document_summary = self.document_summary.get("summary", ""),
                                document_category = self.document_category
                                )
                        }]
                    response = _call_llm(self.llm_provider, messages)
                    if isinstance(response, dict) and "content" in response:
                        translated_text = response["content"]
                        logging.info(f"Current text after reviser (quality): {translated_text[:100]}...")
                    else:
                            translated_text = ""
                    ###############################################################################################
                    word_count_chunk = len(nontranslated.split())
                    word_count_response = len(translated_text.split())
                    
                    if word_count_response < word_count_chunk:
                            logging.warning(f"Translation is too short ({word_count_response} words). Requesting expansion...")
                            # Request expansion if too short
                            expansion_messages = [{
                                "content": (
                                    f"""Please expand this literary translation to match the original length of {word_count_chunk} words while:
                                    - Preserving the author's distinct writing style and voice
                                    - Maintaining consistency with the overall narrative tone and mood
                                    - Elaborating on literary devices, metaphors, and imagery already present
                                    - Using natural sentence structures in the target language
                                    - Avoiding unnecessary padding or repetition
                                    - Ensuring cultural nuances and context are properly conveyed
                                    
                                    Original translation: {translated_text}"""
                                ),
                                "system_message": PROMPT_REVISE_TRANSLATION_SYSTEM.format(
                                target_lang=self.target_language,
                                document_summary = self.document_summary.get("summary", ""),
                                document_category = self.document_category
                                )
                            }]
                            expansion_response = _call_llm(self.llm_provider, expansion_messages)
                            if isinstance(expansion_response, dict) and "content" in expansion_response:
                                translated_text = expansion_response["content"]
                                logging.info(f"Current text after reviser (quality) expanded: {translated_text[:100]}...")
                            else:
                                    translated_text = ""



            return translated_text

def read_file(file_path: str) -> List[Dict]:
    """Reads text from a file, handling PDF and DOCX, preserving structure."""
    if file_path.lower().endswith(".pdf"):
        try:
            elements = partition(filename=file_path)
            structured_text = []
            for el in elements:
                element_type = getattr(el, "type", "Other") # Get the type attribute, or default to "Other"
                if element_type == "Title":
                    structured_text.append({"type": "Title", "content": str(el.text)})
                elif element_type == "NarrativeText":
                    structured_text.append({"type": "Paragraph", "content": str(el.text)})
                else:
                    structured_text.append({"type": "Other", "content": str(el.text)})

            return structured_text
        except Exception as e:
            logging.error(f"Error reading PDF file: {e}")
            return []  # Return empty list on error
    elif file_path.lower().endswith(".docx"):
        try:
            doc = docx.Document(file_path)
            structured_text = []
            for para in doc.paragraphs:
                if para.style.name.startswith("Heading"):
                    structured_text.append({"type": "Title", "content": para.text, "level": int(para.style.name[7:]) if para.style.name[7:].isdigit() else 1})
                else:
                    structured_text.append({"type": "Paragraph", "content": para.text})
            return structured_text
        except Exception as e:
            logging.error(f"Error reading DOCX file: {e}")
            return []  # Return empty list on error
    else:
        try:
            with open(file_path, "r", encoding="utf8") as file:
                return [{"type": "Paragraph", "content": file.read()}]
        except Exception as e:
            logging.error(f"Error reading text file: {e}")
            return []  #Return empty list on error

def write_docx(file_path: str, structured_text: List[Dict], summary: dict = None, category: str = None):
    """Writes structured text to a DOCX file."""
    document = docx.Document()
    for element in structured_text:
        try:
            if element["type"] == "Title":
                level = element.get("level", 1)
                document.add_heading(element["content"], level=level)
            elif element["type"] == "Paragraph":
                document.add_paragraph(element["content"])
            elif element["type"] == "Other":
                document.add_paragraph(element["content"])
            elif element["type"] == "Error":
                document.add_paragraph(f"Error: {element['content']}")
        except Exception as e:
            logging.error(f"Error writing element to docx: {e} - Element: {element}")
            document.add_paragraph(f"Error writing element to docx: {e} - Element: {element}")

    # Add summary and category at the end
    if summary:
      document.add_heading("Document Summary", level=2)
      document.add_paragraph(summary.get("summary", "No summary found"))
      document.add_heading("Document Keywords", level=3)
      document.add_paragraph(", ".join(summary.get("keywords", [])))

    if category:
        document.add_heading("Document Category", level=2)
        document.add_paragraph(category)

    try:
        document.save(file_path)
        logging.info(f"Document saved to: {file_path}")
    except Exception as e:
        logging.error(f"Error saving docx file: {e}")

def main():

    # Initialize the OpenAI class here
    openai_api_key = 'ollama'
    base_url = 'https://b125-34-143-154-113.ngrok-free.app'
    provider = OpenAI(api_key=openai_api_key, model="deepseek-r1:32b-qwen-distill-q8_0", base_url=base_url)

    provider_wrapper = LLMProviderWrapper(provider)
    output_file = 'output_deepseek_32.docx' # Changed extension

    # Initialize the translator with output file
    translator = UniversalTranslator(
        source_language="English",
        target_language="Turkish",
        llm_provider=provider_wrapper,
        country="Turkey",
        max_concurrent_requests=5,
        output_file = output_file,
        max_tokens=1024,
        agent_pipeline=["translator", "improver", "reviser_improver", "quality_checker", "reviser_quality"],
         title_agent_pipeline = ["initial_translator_title", "quality_checker_title", "reviser_title"],
        min_tokens_for_full_pipeline=100, # Set the minimum token length for the full pipeline here
        title_post_process = lambda text: text.upper() #Example of post processing

    )

    # Load and translate the text
    script_dir = ''
    rel_path_to_text = "Surrender_No_preface.docx" # Or change to docx


    logging.info(f"Text to translate: {rel_path_to_text}...")
    logging.info("--------------------------------------------------")

    # Translate the text
    start_time = time.time()
    try:
        translated_elements = translator.translate(os.path.join(script_dir, rel_path_to_text)) #Changed variable name
        end_time = time.time()
        logging.info(f"Time taken: {end_time - start_time} seconds")
    # logging.info(f"Usage: {translator.usage}") #No usage in autogen
        logging.info(f"Translation complete. Check {output_file} for results.")
        logging.debug(f"Translated Elements before writing to file: {translated_elements}") #Log translated elements
        write_docx(output_file, translated_elements, translator.document_summary, translator.document_category) # Call the writer with translated elements and also summary and category
    except Exception as e:
        logging.error(f"Translation failed: {e}")

if __name__ == "__main__":
    main()